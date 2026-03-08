"""Document processing service: extracts text from various file formats."""

from pathlib import Path


async def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    import pdfplumber
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            text += f"\n--- Page {i + 1} ---\n{page_text}"
    return text


async def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    from docx import Document
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            text += row_text + "\n"
    return text


async def extract_data_from_xlsx(file_path: str) -> list[dict]:
    """Extract data from an Excel file as list of dicts."""
    from openpyxl import load_workbook
    wb = load_workbook(file_path, read_only=True, data_only=True)
    all_data = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(h) if h else f"col_{i}" for i, h in enumerate(rows[0])]
        for row in rows[1:]:
            row_dict = {}
            for j, val in enumerate(row):
                if j < len(headers):
                    row_dict[headers[j]] = val
            row_dict["_sheet"] = sheet_name
            all_data.append(row_dict)
    return all_data


async def extract_text_from_image(file_path: str) -> str:
    """For images, we'll use OpenAI vision API to extract text/content."""
    import base64
    with open(file_path, "rb") as f:
        image_data = base64.b64encode(f.read()).decode("utf-8")

    from openai import AsyncOpenAI

    from app.config import settings
    client = AsyncOpenAI(api_key=settings.openai_api_key)

    ext = Path(file_path).suffix.lower()
    mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".gif": "image/gif", ".webp": "image/webp"}
    mime = mime_map.get(ext, "image/png")

    response = await client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": "Extract all text and data from this image. If it contains tables, format them. If it contains charts, describe the data shown."},
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{image_data}"}}
            ]
        }],
        max_tokens=4000,
    )
    return response.choices[0].message.content


async def process_document(file_path: str, filename: str) -> dict:
    """Process a document and return extracted content."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        text = await extract_text_from_pdf(file_path)
        return {"type": "text", "content": text}
    elif ext in (".doc", ".docx"):
        text = await extract_text_from_docx(file_path)
        return {"type": "text", "content": text}
    elif ext in (".xls", ".xlsx"):
        data = await extract_data_from_xlsx(file_path)
        return {"type": "spreadsheet", "content": data}
    elif ext in (".png", ".jpg", ".jpeg", ".gif", ".webp"):
        text = await extract_text_from_image(file_path)
        return {"type": "text", "content": text}
    elif ext in (".txt", ".csv", ".md"):
        with open(file_path, errors="replace") as f:
            text = f.read()
        return {"type": "text", "content": text}
    else:
        return {"type": "unknown", "content": f"Unsupported file type: {ext}"}
