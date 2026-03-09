#!/usr/bin/env python3
"""Generate realistic sample evidence documents for Lincoln Innovation Academy.

Run: python generate_sample_data.py
Output: sample_data/ directory with 30+ evidence files covering all 9 SQF dimensions
"""

import os
from pathlib import Path
from datetime import date

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
    ListFlowable, ListItem,
)

OUT_DIR = Path(__file__).parent / "sample_data"

HEADER_FILL = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
BOLD_FONT = Font(bold=True, size=11)
THIN_BORDER = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"), bottom=Side(style="thin"),
)


def style_header_row(ws, row, cols):
    for c in range(1, cols + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
        cell.border = THIN_BORDER


def auto_width(ws):
    for col in ws.columns:
        max_len = max((len(str(cell.value or "")) for cell in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)


# ─── 1. Teacher Retention Analysis (XLSX) ───────────────────────────

def generate_teacher_retention():
    wb = Workbook()

    # Sheet 1: Retention by Department
    ws = wb.active
    ws.title = "Retention by Department"
    headers = ["Department", "2021-22 Staff", "2021-22 Retained", "2021-22 Rate",
               "2022-23 Staff", "2022-23 Retained", "2022-23 Rate",
               "2023-24 Staff", "2023-24 Retained", "2023-24 Rate"]
    ws.append(headers)
    style_header_row(ws, 1, len(headers))

    depts = [
        ("ELA", 8, 6, 9, 7, 9, 8),
        ("Math", 7, 5, 7, 6, 8, 7),
        ("Science", 5, 4, 5, 4, 6, 5),
        ("Social Studies", 4, 3, 5, 4, 5, 4),
        ("Special Education", 6, 4, 6, 5, 7, 6),
        ("Specials (Art/Music/PE)", 5, 4, 5, 4, 5, 5),
        ("Administration", 4, 4, 4, 4, 4, 4),
        ("Support Staff", 6, 4, 6, 5, 7, 6),
    ]
    for dept, s1, r1, s2, r2, s3, r3 in depts:
        ws.append([dept, s1, r1, f"{r1/s1:.0%}", s2, r2, f"{r2/s2:.0%}", s3, r3, f"{r3/s3:.0%}"])

    totals = [sum(d[i] for d in depts) for i in range(1, 7)]
    ws.append(["TOTAL", totals[0], totals[1], f"{totals[1]/totals[0]:.0%}",
               totals[2], totals[3], f"{totals[3]/totals[2]:.0%}",
               totals[4], totals[5], f"{totals[4] and totals[5]/totals[4]:.0%}"])
    for cell in ws[ws.max_row]:
        cell.font = BOLD_FONT
    auto_width(ws)

    # Sheet 2: Demographic Breakdown
    ws2 = wb.create_sheet("Demographic Breakdown")
    ws2.append(["Demographic Group", "Total Staff", "Retained", "Rate", "District Avg Rate"])
    style_header_row(ws2, 1, 5)
    demos = [
        ("White", 28, 24, "86%", "88%"),
        ("Black/African American", 9, 7, "78%", "80%"),
        ("Hispanic/Latino", 8, 6, "75%", "79%"),
        ("Asian", 3, 3, "100%", "90%"),
        ("Two or More Races", 3, 2, "67%", "82%"),
        ("Male", 18, 15, "83%", "84%"),
        ("Female", 33, 27, "82%", "86%"),
        ("0-3 Years Experience", 15, 10, "67%", "70%"),
        ("4-10 Years Experience", 22, 19, "86%", "85%"),
        ("11+ Years Experience", 14, 13, "93%", "92%"),
    ]
    for row in demos:
        ws2.append(row)
    auto_width(ws2)

    # Sheet 3: Exit Interview Themes
    ws3 = wb.create_sheet("Exit Interview Themes")
    ws3.append(["Theme", "Mentions (2021-22)", "Mentions (2022-23)", "Mentions (2023-24)", "Trend"])
    style_header_row(ws3, 1, 5)
    themes = [
        ("Compensation concerns", 8, 6, 4, "Improving"),
        ("Lack of admin support", 5, 4, 2, "Improving"),
        ("Work-life balance", 6, 5, 5, "Stable"),
        ("Limited growth opportunities", 4, 4, 3, "Slight improvement"),
        ("Student behavior challenges", 7, 5, 3, "Improving"),
        ("Relocation/personal reasons", 3, 3, 4, "Stable"),
        ("Better opportunity elsewhere", 4, 3, 2, "Improving"),
        ("School culture/climate", 3, 2, 1, "Improving"),
    ]
    for row in themes:
        ws3.append(row)
    auto_width(ws3)

    wb.save(OUT_DIR / "teacher_retention_analysis.xlsx")
    print("  ✓ teacher_retention_analysis.xlsx")


# ─── 2. Classroom Observation Report (PDF) ──────────────────────────

def generate_classroom_observation():
    path = OUT_DIR / "classroom_observation_report.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("CustomTitle", parent=styles["Title"], fontSize=18,
                                  textColor=colors.HexColor("#1a365d"), spaceAfter=6)
    subtitle_style = ParagraphStyle("Subtitle", parent=styles["Normal"], fontSize=12,
                                     textColor=colors.HexColor("#4a5568"), spaceAfter=20)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14,
                         textColor=colors.HexColor("#2d3748"), spaceBefore=16, spaceAfter=8)
    body = styles["Normal"]
    body.fontSize = 10
    body.leading = 14

    story = []
    story.append(Paragraph("Lincoln Innovation Academy", title_style))
    story.append(Paragraph("Classroom Observation Summary Report — 2023-24 School Year", subtitle_style))
    story.append(Paragraph("Prepared by: Meridian Consulting Group | Date: March 2024", body))
    story.append(Spacer(1, 20))

    story.append(Paragraph("1. Overview", h2))
    story.append(Paragraph(
        "Between October 2023 and February 2024, the consulting team conducted 47 classroom observations "
        "across all grade levels (K-8) at Lincoln Innovation Academy. Observations used the Instructional "
        "Practice Guide (IPG) framework and lasted 20-30 minutes each. This report summarizes key findings, "
        "patterns of strength, and areas for growth.", body))
    story.append(Spacer(1, 12))

    # Observation counts table
    story.append(Paragraph("2. Observation Distribution", h2))
    obs_data = [
        ["Grade Band", "Observations", "ELA", "Math", "Science", "Other"],
        ["K-2", "16", "7", "5", "2", "2"],
        ["3-5", "18", "7", "6", "3", "2"],
        ["6-8", "13", "5", "4", "2", "2"],
        ["Total", "47", "19", "15", "7", "6"],
    ]
    t = Table(obs_data, colWidths=[1.5*inch, 1*inch, 0.8*inch, 0.8*inch, 0.8*inch, 0.8*inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#f0f4f8")]),
    ]))
    story.append(t)
    story.append(Spacer(1, 16))

    # Ratings table
    story.append(Paragraph("3. Instructional Practice Ratings", h2))
    story.append(Paragraph(
        "Each observation was rated on four IPG dimensions using a 1-4 scale "
        "(1=Not Yet, 2=Developing, 3=Proficient, 4=Distinguished).", body))
    story.append(Spacer(1, 8))

    ratings_data = [
        ["Dimension", "K-2 Avg", "3-5 Avg", "6-8 Avg", "School Avg"],
        ["Core Content Alignment", "2.9", "2.7", "2.5", "2.7"],
        ["Student Engagement", "3.1", "2.8", "2.4", "2.8"],
        ["Differentiation & Support", "2.4", "2.3", "2.1", "2.3"],
        ["Checks for Understanding", "2.6", "2.5", "2.3", "2.5"],
    ]
    t2 = Table(ratings_data, colWidths=[2*inch, 1*inch, 1*inch, 1*inch, 1*inch])
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
    ]))
    story.append(t2)
    story.append(Spacer(1, 16))

    story.append(PageBreak())

    # Strengths
    story.append(Paragraph("4. Key Strengths", h2))
    strengths = [
        "Strong classroom routines and procedures in K-2 classrooms, with consistent use of "
        "transition signals and clear behavioral expectations.",
        "High levels of student-to-student discourse in ELA classes, particularly in grades 3-5 "
        "where teachers regularly use structured discussion protocols (Think-Pair-Share, Socratic Seminars).",
        "Consistent posting and referencing of learning objectives across most classrooms. 89% of "
        "observed classrooms had clear, standards-aligned objectives posted.",
        "Positive teacher-student relationships evident across grade levels. Teachers consistently "
        "used warm, respectful language and demonstrated genuine care for student well-being.",
        "Effective use of anchor charts and visual supports, especially in K-2 and special education "
        "classrooms, providing students with accessible reference materials.",
    ]
    items = [ListItem(Paragraph(s, body)) for s in strengths]
    story.append(ListFlowable(items, bulletType="bullet", start="•"))
    story.append(Spacer(1, 16))

    # Growth areas
    story.append(Paragraph("5. Areas for Growth", h2))
    growth = [
        "Differentiation and scaffolding: Only 34% of observed lessons included meaningful "
        "differentiation for diverse learners. Middle school classrooms showed the greatest need, "
        "with most lessons following a whole-group, one-size-fits-all approach.",
        "Rigor of questioning: While teachers asked frequent questions, only 28% of questions "
        "observed were at DOK 3 or above. Most questioning remained at the recall/comprehension level, "
        "limiting opportunities for critical thinking.",
        "Formative assessment and responsive teaching: Teachers infrequently adjusted instruction "
        "based on real-time student understanding. Exit tickets were used in only 40% of observed "
        "lessons, and when used, rarely informed next-day instruction.",
        "Student agency and metacognition: Students in 6-8 classrooms had limited opportunities "
        "to set goals, self-assess, or make choices about their learning pathways.",
        "Technology integration: While devices were available, technology use was primarily "
        "consumptive (watching videos, completing online worksheets) rather than for creation, "
        "collaboration, or deeper learning.",
    ]
    items = [ListItem(Paragraph(s, body)) for s in growth]
    story.append(ListFlowable(items, bulletType="bullet", start="•"))
    story.append(Spacer(1, 16))

    # Recommendations
    story.append(Paragraph("6. Recommendations", h2))
    recs = [
        "Implement school-wide differentiation framework with coaching support, prioritizing "
        "math and science in grades 6-8.",
        "Develop a questioning toolkit for teachers with DOK-leveled question stems by content area.",
        "Establish a formative assessment cycle: exit ticket → data review → responsive planning. "
        "Provide dedicated PD and planning time.",
        "Launch student-led conferencing initiative in grades 3-8 to build metacognitive skills.",
        "Create a technology integration rubric aligned to SAMR model and provide targeted PD.",
    ]
    items = [ListItem(Paragraph(s, body)) for s in recs]
    story.append(ListFlowable(items, bulletType="bullet", start="•"))

    doc.build(story)
    print("  ✓ classroom_observation_report.pdf")


# ─── 3. Student Achievement Data (XLSX) ─────────────────────────────

def generate_student_achievement():
    wb = Workbook()

    # Sheet 1: State Test Proficiency
    ws = wb.active
    ws.title = "State Test Proficiency"
    ws.append(["Grade", "ELA 2022", "ELA 2023", "ELA 2024", "Math 2022", "Math 2023", "Math 2024",
               "Science 2023", "Science 2024"])
    style_header_row(ws, 1, 9)
    data = [
        ("3", 42, 45, 49, 38, 41, 44, None, None),
        ("4", 44, 48, 52, 40, 43, 46, None, None),
        ("5", 46, 49, 53, 37, 40, 45, 41, 44),
        ("6", 41, 44, 48, 35, 38, 42, 38, 42),
        ("7", 39, 42, 46, 33, 36, 40, 36, 39),
        ("8", 38, 41, 45, 31, 34, 38, 35, 38),
        ("School Avg", 42, 45, 49, 36, 39, 43, 38, 41),
        ("District Avg", 48, 49, 50, 44, 45, 46, 44, 45),
        ("State Avg", 51, 52, 52, 46, 47, 47, 46, 47),
    ]
    for row in data:
        ws.append([row[0]] + [f"{v}%" if v else "N/A" for v in row[1:]])
    for cell in ws[ws.max_row - 1]:
        cell.font = BOLD_FONT
    auto_width(ws)

    # Sheet 2: Growth Percentiles
    ws2 = wb.create_sheet("Growth Percentiles")
    ws2.append(["Grade", "ELA MGP 2023", "ELA MGP 2024", "Math MGP 2023", "Math MGP 2024"])
    style_header_row(ws2, 1, 5)
    growth = [
        ("3", 52, 56, 48, 53),
        ("4", 54, 58, 50, 55),
        ("5", 51, 55, 47, 52),
        ("6", 49, 53, 45, 50),
        ("7", 48, 51, 44, 48),
        ("8", 47, 50, 43, 47),
        ("School Avg", 50, 54, 46, 51),
    ]
    for row in growth:
        ws2.append(row)
    ws2.append([])
    ws2.append(["Note: MGP = Median Growth Percentile. 50 = typical growth. >55 = high growth."])
    auto_width(ws2)

    # Sheet 3: Subgroup Performance
    ws3 = wb.create_sheet("Subgroup Performance")
    ws3.append(["Subgroup", "N", "ELA Prof 2024", "Math Prof 2024", "ELA Growth 2024", "Math Growth 2024"])
    style_header_row(ws3, 1, 6)
    subs = [
        ("All Students", 412, "49%", "43%", 54, 51),
        ("White", 165, "58%", "52%", 55, 53),
        ("Black/African American", 103, "39%", "33%", 52, 49),
        ("Hispanic/Latino", 99, "42%", "36%", 54, 50),
        ("Asian", 25, "64%", "60%", 57, 56),
        ("Economically Disadvantaged", 247, "38%", "32%", 52, 48),
        ("English Learners", 58, "28%", "25%", 50, 46),
        ("Students with Disabilities", 62, "22%", "18%", 48, 44),
    ]
    for row in subs:
        ws3.append(row)
    auto_width(ws3)

    wb.save(OUT_DIR / "student_achievement_data.xlsx")
    print("  ✓ student_achievement_data.xlsx")


# ─── 4. Family Survey Results (PDF) ─────────────────────────────────

def generate_family_survey():
    path = OUT_DIR / "family_survey_results.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("CustomTitle", parent=styles["Title"], fontSize=18,
                                  textColor=colors.HexColor("#1a365d"), spaceAfter=6)
    subtitle_style = ParagraphStyle("Subtitle", parent=styles["Normal"], fontSize=12,
                                     textColor=colors.HexColor("#4a5568"), spaceAfter=20)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14,
                         textColor=colors.HexColor("#2d3748"), spaceBefore=16, spaceAfter=8)
    body = styles["Normal"]
    body.fontSize = 10
    body.leading = 14

    story = []
    story.append(Paragraph("Lincoln Innovation Academy", title_style))
    story.append(Paragraph("Annual Family Survey Results — Spring 2024", subtitle_style))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Response Summary", h2))
    story.append(Paragraph(
        "The annual family survey was administered in March 2024 via online form and paper copies "
        "sent home with students. A total of 287 families responded out of 412 enrolled students, "
        "yielding a <b>69.7% response rate</b> (up from 58% in 2023). Surveys were available in "
        "English and Spanish.", body))
    story.append(Spacer(1, 16))

    # Satisfaction ratings
    story.append(Paragraph("Satisfaction Ratings by Domain", h2))
    story.append(Paragraph(
        "Families rated their agreement on a 5-point scale (1=Strongly Disagree to 5=Strongly Agree). "
        "Results below show the percentage of respondents selecting Agree or Strongly Agree.", body))
    story.append(Spacer(1, 8))

    survey_data = [
        ["Domain", "2023 %Agree", "2024 %Agree", "Change"],
        ["My child is safe at school", "74%", "81%", "+7"],
        ["The school has high academic expectations", "68%", "76%", "+8"],
        ["Teachers communicate regularly about my child's progress", "61%", "72%", "+11"],
        ["My child feels they belong at this school", "77%", "83%", "+6"],
        ["The school welcomes family involvement", "65%", "74%", "+9"],
        ["I am satisfied with school leadership", "62%", "71%", "+9"],
        ["My child is challenged academically", "58%", "67%", "+9"],
        ["The school addresses bullying effectively", "52%", "63%", "+11"],
        ["I would recommend this school to other families", "66%", "75%", "+9"],
    ]
    t = Table(survey_data, colWidths=[2.8*inch, 1.1*inch, 1.1*inch, 0.8*inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
    ]))
    story.append(t)
    story.append(Spacer(1, 16))

    story.append(PageBreak())

    # Open-ended themes
    story.append(Paragraph("Open-Ended Response Themes", h2))
    story.append(Paragraph(
        "217 families provided open-ended comments. Responses were coded into themes:", body))
    story.append(Spacer(1, 8))

    themes_data = [
        ["Theme", "Mentions", "Sentiment"],
        ["Teacher quality and dedication", "89", "Positive"],
        ["Desire for more communication", "67", "Mixed"],
        ["After-school program requests", "54", "Neutral"],
        ["Homework concerns (too much/too little)", "48", "Negative"],
        ["Appreciation for school culture improvements", "45", "Positive"],
        ["Concerns about class sizes", "38", "Negative"],
        ["Request for more STEM programming", "34", "Neutral"],
        ["Transportation issues", "28", "Negative"],
        ["Praise for new leadership team", "26", "Positive"],
    ]
    t2 = Table(themes_data, colWidths=[2.8*inch, 1*inch, 1*inch])
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
    ]))
    story.append(t2)
    story.append(Spacer(1, 16))

    # Selected comments
    story.append(Paragraph("Selected Family Comments", h2))
    comments = [
        '"The teachers really care about my kids. Ms. Rodriguez goes above and beyond every single day."',
        '"I wish there was more communication about what my child is learning. I only hear about problems."',
        '"The new principal has made a huge difference. The school feels safer and more organized this year."',
        '"My son needs more challenge — he finishes work early and just sits there. Need gifted programming."',
        '"We love Lincoln but the after-school options are limited compared to other schools in the district."',
        '"Math instruction has improved a lot since they adopted the new curriculum. My daughter actually enjoys math now."',
    ]
    for c in comments:
        story.append(Paragraph(f"<i>{c}</i>", body))
        story.append(Spacer(1, 6))

    doc.build(story)
    print("  ✓ family_survey_results.pdf")


# ─── 5. PD Calendar (DOCX) ──────────────────────────────────────────

def generate_pd_calendar():
    doc = Document()

    # Title
    title = doc.add_heading("Lincoln Innovation Academy", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    doc.add_heading("Professional Development Calendar — 2023-24", level=1)
    doc.add_paragraph(
        "This calendar outlines the school's professional development plan for the 2023-24 academic year. "
        "All sessions are mandatory unless marked optional. Total PD hours: 72 hours."
    )

    # PD sessions
    sessions = [
        ("Aug 14-18, 2023", "Opening Institute", "Full Staff", "Dr. Maria Santos, Principal",
         "Mission & vision review, SQF framework introduction, data review, "
         "classroom setup, team planning", "30 hrs"),
        ("Sep 15, 2023", "Data-Driven Instruction", "All Teachers", "Dr. James Morton (external)",
         "Assessment literacy, interim assessment analysis, action planning protocols, "
         "student work analysis", "6 hrs"),
        ("Oct 20, 2023", "Differentiated Instruction I", "All Teachers", "Linda Park, Instructional Coach",
         "UDL framework introduction, flexible grouping strategies, tiered assignments, "
         "scaffolding techniques", "6 hrs"),
        ("Nov 17, 2023", "Culturally Responsive Teaching", "All Teachers", "Dr. Keisha Williams (external)",
         "Cultural competence self-assessment, asset-based framing, "
         "culturally sustaining pedagogy practices", "6 hrs"),
        ("Jan 12, 2024", "Mid-Year Data Review", "All Teachers", "Dr. Maria Santos & Linda Park",
         "Interim assessment results analysis, progress monitoring, goal recalibration, "
         "intervention planning", "6 hrs"),
        ("Feb 9, 2024", "Differentiated Instruction II", "All Teachers", "Linda Park, Instructional Coach",
         "Advanced differentiation strategies, technology-enhanced differentiation, "
         "co-teaching models, peer observation debrief", "6 hrs"),
        ("Mar 15, 2024", "Social-Emotional Learning", "All Staff", "School Counseling Team",
         "SEL competencies integration, trauma-informed practices, restorative circles, "
         "de-escalation strategies", "6 hrs"),
        ("Apr 12, 2024", "Standards-Based Grading (Optional)", "Interested Teachers",
         "Dr. James Morton (external)",
         "Proficiency scales, evidence-based grading, student self-assessment, "
         "gradebook alignment", "3 hrs"),
        ("May 10, 2024", "End-of-Year Reflection", "Full Staff", "Leadership Team",
         "Year in review, SQF self-assessment, summer planning, celebrations", "3 hrs"),
    ]

    table = doc.add_table(rows=1, cols=6, style="Light Grid Accent 1")
    hdr = table.rows[0].cells
    for i, text in enumerate(["Date", "Session Title", "Audience", "Facilitator", "Topics", "Hours"]):
        hdr[i].text = text
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    for date_str, title, audience, facilitator, topics, hours in sessions:
        row_cells = table.add_row().cells
        row_cells[0].text = date_str
        row_cells[1].text = title
        row_cells[2].text = audience
        row_cells[3].text = facilitator
        row_cells[4].text = topics
        row_cells[5].text = hours

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph()
    doc.add_heading("Ongoing Professional Learning", level=2)
    ongoing = [
        "Weekly grade-level team meetings (45 min) — collaborative planning and data review",
        "Monthly instructional rounds — peer observation with structured debrief protocol",
        "Bi-weekly coaching cycles for teachers in years 1-3 with instructional coach",
        "Optional book study: 'Culturally Responsive Teaching and the Brain' by Zaretta Hammond",
        "New teacher mentor program — monthly mentor-mentee meetings and quarterly cohort sessions",
    ]
    for item in ongoing:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(str(OUT_DIR / "pd_calendar_2024.docx"))
    print("  ✓ pd_calendar_2024.docx")


# ─── 6. Budget Summary (XLSX) ───────────────────────────────────────

def generate_budget():
    wb = Workbook()

    # Sheet 1: Revenue
    ws = wb.active
    ws.title = "Revenue"
    ws.append(["Revenue Source", "FY2023 Budget", "FY2023 Actual", "FY2024 Budget", "% Change"])
    style_header_row(ws, 1, 5)
    rev = [
        ("Per-Pupil State Funding", 3_280_000, 3_280_000, 3_460_000, "5.5%"),
        ("Title I Federal Funds", 420_000, 418_500, 435_000, "3.6%"),
        ("Title II (Teacher Quality)", 85_000, 85_000, 88_000, "3.5%"),
        ("IDEA Special Education", 310_000, 308_200, 325_000, "4.8%"),
        ("State Special Education", 180_000, 180_000, 192_000, "6.7%"),
        ("Local Grants & Donations", 95_000, 112_400, 105_000, "10.5%"),
        ("E-Rate Technology", 45_000, 44_800, 48_000, "6.7%"),
        ("Food Service Revenue", 210_000, 205_600, 218_000, "3.8%"),
    ]
    for name, b23, a23, b24, pct in rev:
        ws.append([name, f"${b23:,}", f"${a23:,}", f"${b24:,}", pct])
    total_b23 = sum(r[1] for r in rev)
    total_a23 = sum(r[2] for r in rev)
    total_b24 = sum(r[3] for r in rev)
    ws.append(["TOTAL REVENUE", f"${total_b23:,}", f"${total_a23:,}", f"${total_b24:,}",
               f"{(total_b24 - total_b23) / total_b23:.1%}"])
    for cell in ws[ws.max_row]:
        cell.font = BOLD_FONT
    auto_width(ws)

    # Sheet 2: Expenditures
    ws2 = wb.create_sheet("Expenditures")
    ws2.append(["Category", "FY2023 Budget", "FY2023 Actual", "FY2024 Budget", "% of Total"])
    style_header_row(ws2, 1, 5)
    exp = [
        ("Teacher Salaries", 2_180_000, 2_165_000, 2_310_000, "47.4%"),
        ("Teacher Benefits", 654_000, 649_500, 693_000, "14.2%"),
        ("Admin Salaries & Benefits", 385_000, 385_000, 398_000, "8.2%"),
        ("Support Staff Salaries", 290_000, 288_000, 305_000, "6.3%"),
        ("Instructional Materials", 185_000, 178_400, 198_000, "4.1%"),
        ("Technology", 145_000, 152_300, 160_000, "3.3%"),
        ("Professional Development", 95_000, 88_600, 110_000, "2.3%"),
        ("Facilities & Maintenance", 280_000, 275_200, 295_000, "6.1%"),
        ("Student Services", 120_000, 118_400, 130_000, "2.7%"),
        ("Food Service", 195_000, 198_100, 210_000, "4.3%"),
        ("Other Operating", 56_000, 52_800, 62_000, "1.3%"),
    ]
    for name, b23, a23, b24, pct in exp:
        ws2.append([name, f"${b23:,}", f"${a23:,}", f"${b24:,}", pct])
    total_eb23 = sum(e[1] for e in exp)
    total_ea23 = sum(e[2] for e in exp)
    total_eb24 = sum(e[3] for e in exp)
    ws2.append(["TOTAL EXPENDITURES", f"${total_eb23:,}", f"${total_ea23:,}", f"${total_eb24:,}", "100%"])
    for cell in ws2[ws2.max_row]:
        cell.font = BOLD_FONT
    auto_width(ws2)

    # Sheet 3: Per-Pupil & Ratios
    ws3 = wb.create_sheet("Per-Pupil Analysis")
    ws3.append(["Metric", "FY2023", "FY2024", "District Avg"])
    style_header_row(ws3, 1, 4)
    metrics = [
        ("Enrollment", "398", "412", "445"),
        ("Total Per-Pupil Spending", "$11,834", "$11,822", "$12,450"),
        ("Instructional Spending Per Pupil", "$7,621", "$7,771", "$8,100"),
        ("Instructional % of Budget", "64.4%", "65.7%", "65.1%"),
        ("Student-to-Teacher Ratio", "14.2:1", "13.7:1", "15.1:1"),
        ("PD Spending Per Teacher", "$3,154", "$3,548", "$2,800"),
        ("Technology Spending Per Student", "$383", "$388", "$350"),
    ]
    for row in metrics:
        ws3.append(row)
    auto_width(ws3)

    wb.save(OUT_DIR / "budget_summary_fy2024.xlsx")
    print("  ✓ budget_summary_fy2024.xlsx")


# ─── 7. Leadership Team Notes (DOCX) ────────────────────────────────

def generate_leadership_notes():
    doc = Document()
    title = doc.add_heading("Lincoln Innovation Academy — Leadership Team Meeting Notes", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    meetings = [
        {
            "date": "January 18, 2024",
            "attendees": "Dr. Maria Santos (Principal), James Wilson (AP), Linda Park (Instructional Coach), "
                         "Rosa Martinez (SPED Coordinator), David Chen (Dean of Culture), "
                         "Keisha Thompson (Grade-Level Chair, K-2), Mark Sullivan (Grade-Level Chair, 3-5), "
                         "Angela Foster (Grade-Level Chair, 6-8)",
            "agenda": [
                ("Mid-Year Assessment Data Review",
                 "Linda presented interim assessment results showing 4-point growth in ELA and 3-point growth "
                 "in Math compared to fall benchmarks. Growth is strongest in grades 3-5. Middle school math "
                 "remains a concern with only 38% of 6-8 students on track for grade-level proficiency. "
                 "Rosa noted that SPED students showed above-average growth in co-taught classrooms."),
                ("Observation Walkthrough Debrief",
                 "The team reviewed classroom observation data from the consulting team's November visit. "
                 "Key discussion points: differentiation gaps across grade levels, need for higher-order "
                 "questioning strategies, and inconsistent use of formative assessment. Linda will develop "
                 "targeted coaching plans for teachers identified as needing support."),
                ("Family Engagement Update",
                 "David reported that the new family communication system (weekly newsletters + Remind app) "
                 "has increased parent engagement metrics by 22%. The January family night had 145 attendees, "
                 "up from 89 at the same event last year. Spring survey distribution plan was reviewed."),
            ],
            "action_items": [
                "Linda: Create differentiation coaching plans for 8 identified teachers by Feb 1",
                "Mark & Angela: Develop grade-band questioning stems document by Feb 9",
                "Rosa: Schedule co-teaching model observations for all SPED teachers by Jan 31",
                "David: Finalize spring survey questions and translation by Feb 5",
                "Dr. Santos: Follow up with district on additional math intervention funding",
            ],
        },
        {
            "date": "February 15, 2024",
            "attendees": "Dr. Maria Santos (Principal), James Wilson (AP), Linda Park (Instructional Coach), "
                         "Rosa Martinez (SPED Coordinator), David Chen (Dean of Culture), "
                         "Keisha Thompson (K-2 Chair), Angela Foster (6-8 Chair). "
                         "Mark Sullivan absent (professional conference).",
            "agenda": [
                ("Coaching Cycle Update",
                 "Linda reported that 6 of 8 targeted teachers have completed their first coaching cycle. "
                 "Initial feedback is positive. Two teachers in grades 6-7 are resistant to observation-based "
                 "feedback. Dr. Santos will have individual conversations with these teachers about the "
                 "supportive intent of coaching."),
                ("Student Culture & Discipline Data",
                 "David presented Q2 discipline data: office referrals down 18% compared to Q1, and 31% lower "
                 "than the same period last year. Suspensions are down 40%. The restorative practices pilot "
                 "in grades 6-8 is showing promising results, though some teachers report feeling undertrained. "
                 "Additional restorative practices PD was requested."),
                ("Staffing and Retention Planning",
                 "Dr. Santos shared early retention data: 3 teachers have submitted non-renewal notices "
                 "(2 relocations, 1 career change). Recruitment for replacements will begin immediately. "
                 "The team discussed strategies to retain high-performing teachers, including a new teacher "
                 "leadership pathway and stipend for department leads."),
            ],
            "action_items": [
                "Dr. Santos: Meet with reluctant teachers individually by Feb 22",
                "David: Organize additional restorative practices training for March PD day",
                "James: Post open positions and begin recruitment pipeline by Feb 23",
                "Dr. Santos & James: Develop teacher leadership pathway proposal by March 1",
                "Keisha: Share K-2 co-teaching model successes at next full staff meeting",
            ],
        },
        {
            "date": "March 14, 2024",
            "attendees": "Full leadership team present.",
            "agenda": [
                ("SQF Self-Assessment Progress",
                 "The team reviewed progress on the School Quality Framework self-assessment. Of 43 components, "
                 "the team has completed self-ratings for 28. Strongest areas: Organizational Purpose (Dimension 1) "
                 "and School Culture (Dimension 6). Areas needing more evidence: Academic Program details "
                 "(Dimension 2) and Talent Management systems (Dimension 5). The consulting team's next visit "
                 "is scheduled for April to review evidence and validate ratings."),
                ("Spring Assessment Preparation",
                 "Mark outlined the state testing schedule (April 8-19) and logistics plan. Testing coordinators "
                 "have been assigned. The team discussed the balance between test prep and maintaining regular "
                 "instruction. Decision: No dedicated test prep blocks; instead, embed test-taking strategies "
                 "into regular instruction."),
                ("Action Plan Development",
                 "The team began drafting the improvement action plan based on assessment findings to date. "
                 "Priority areas identified: (1) Differentiated instruction, (2) Data-driven decision making, "
                 "(3) Family engagement systems, (4) Teacher retention and development. Each leadership team "
                 "member will own one priority area."),
            ],
            "action_items": [
                "All chairs: Complete remaining SQF component self-ratings by March 28",
                "Linda: Compile evidence portfolio for Dimensions 2 and 5 by March 25",
                "Mark: Distribute testing logistics plan to all teachers by March 18",
                "Team leads: Draft 90-day action plans for assigned priority areas by April 1",
                "Dr. Santos: Confirm consulting team visit dates for April",
            ],
        },
    ]

    for meeting in meetings:
        doc.add_heading(f"Meeting Date: {meeting['date']}", level=1)
        doc.add_paragraph(f"Attendees: {meeting['attendees']}")
        doc.add_heading("Agenda & Discussion", level=2)

        for topic, discussion in meeting["agenda"]:
            doc.add_heading(topic, level=3)
            doc.add_paragraph(discussion)

        doc.add_heading("Action Items", level=2)
        for item in meeting["action_items"]:
            doc.add_paragraph(item, style="List Bullet")

        doc.add_paragraph()  # spacer
        doc.add_paragraph("―" * 60)

    doc.save(str(OUT_DIR / "leadership_team_notes.docx"))
    print("  ✓ leadership_team_notes.docx")


# ─── 8. Student Culture Walkthrough (PDF) ────────────────────────────

def generate_culture_walkthrough():
    path = OUT_DIR / "student_culture_walkthrough.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("CustomTitle", parent=styles["Title"], fontSize=18,
                                  textColor=colors.HexColor("#1a365d"), spaceAfter=6)
    subtitle_style = ParagraphStyle("Subtitle", parent=styles["Normal"], fontSize=12,
                                     textColor=colors.HexColor("#4a5568"), spaceAfter=20)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14,
                         textColor=colors.HexColor("#2d3748"), spaceBefore=16, spaceAfter=8)
    body = styles["Normal"]
    body.fontSize = 10
    body.leading = 14

    story = []
    story.append(Paragraph("Lincoln Innovation Academy", title_style))
    story.append(Paragraph("School Culture & Climate Walkthrough Report — February 2024", subtitle_style))
    story.append(Paragraph("Observers: Sarah Chen (Lead Consultant), David Park (Associate)", body))
    story.append(Spacer(1, 20))

    story.append(Paragraph("1. Walkthrough Summary", h2))
    story.append(Paragraph(
        "On February 8, 2024, consultants conducted a full-day culture and climate walkthrough at Lincoln "
        "Innovation Academy. The walkthrough included observations of hallway transitions (4 transition "
        "periods), classroom environments (12 classrooms visited), common areas (cafeteria, gym, "
        "library), and conversations with 15 students, 8 teachers, and 3 support staff members.", body))
    story.append(Spacer(1, 12))

    # Environment observations
    story.append(Paragraph("2. Physical Environment", h2))
    env_data = [
        ["Indicator", "K-2 Wing", "3-5 Wing", "6-8 Wing", "Common Areas"],
        ["Hallways clean and well-maintained", "Yes", "Yes", "Mostly", "Yes"],
        ["Student work displayed", "Extensive", "Good", "Minimal", "Some"],
        ["Mission/values visibly posted", "Yes", "Yes", "Yes", "Yes"],
        ["Welcoming entrance/signage", "—", "—", "—", "Good"],
        ["Library accessible and inviting", "—", "—", "—", "Excellent"],
    ]
    t = Table(env_data, colWidths=[2*inch, 1*inch, 1*inch, 1*inch, 1.2*inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    story.append(t)
    story.append(Spacer(1, 16))

    # Transitions
    story.append(Paragraph("3. Hallway Transitions", h2))
    story.append(Paragraph(
        "Four transition periods were observed (8:15 AM arrival, 10:30 AM specials rotation, "
        "12:00 PM lunch, 2:45 PM dismissal):", body))
    story.append(Spacer(1, 8))
    transitions = [
        "<b>Morning arrival (8:15 AM):</b> Well-organized. Staff stationed at key points greeting "
        "students by name. Breakfast available in cafeteria. Students transitioned to classrooms "
        "within 8 minutes. Warm, positive tone throughout.",
        "<b>Specials rotation (10:30 AM):</b> K-2 transitions were smooth with teachers escorting "
        "classes. Grades 3-5 showed some hallway congestion but students were respectful. "
        "6-8 transitions had 3-4 students lingering without urgency; hall monitors redirected effectively.",
        "<b>Lunch (12:00 PM):</b> Cafeteria was orderly. Noise levels were appropriate. Students "
        "cleaned up after themselves. Lunch aides were visible and engaged. One minor conflict between "
        "two 7th-grade students was de-escalated quickly by the Dean of Culture using restorative language.",
        "<b>Dismissal (2:45 PM):</b> Organized system with clear traffic patterns. Car riders, walkers, "
        "and bus riders separated efficiently. After-school program students reported directly to "
        "designated rooms.",
    ]
    items = [ListItem(Paragraph(s, body)) for s in transitions]
    story.append(ListFlowable(items, bulletType="bullet", start="•"))

    story.append(PageBreak())

    # Classroom culture
    story.append(Paragraph("4. Classroom Culture Indicators", h2))
    culture_data = [
        ["Indicator", "Observed In", "% of Rooms"],
        ["Posted classroom expectations/norms", "11/12", "92%"],
        ["Positive teacher language and tone", "10/12", "83%"],
        ["Student voice and choice opportunities", "6/12", "50%"],
        ["Restorative practices visible (circles, check-ins)", "5/12", "42%"],
        ["Growth mindset language/displays", "8/12", "67%"],
        ["Inclusive representation in materials", "7/12", "58%"],
        ["Students engaged and on-task", "9/12", "75%"],
        ["Collaborative learning structures", "7/12", "58%"],
    ]
    t2 = Table(culture_data, colWidths=[2.8*inch, 1*inch, 1*inch])
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
    ]))
    story.append(t2)
    story.append(Spacer(1, 16))

    # Discipline data
    story.append(Paragraph("5. Discipline Data Summary (Year-to-Date)", h2))
    disc_data = [
        ["Metric", "2022-23 YTD", "2023-24 YTD", "Change"],
        ["Office Referrals", "187", "124", "-34%"],
        ["In-School Suspensions", "42", "28", "-33%"],
        ["Out-of-School Suspensions", "18", "11", "-39%"],
        ["Chronic Absenteeism Rate", "14.2%", "11.8%", "-2.4 pts"],
        ["Average Daily Attendance", "93.1%", "94.6%", "+1.5 pts"],
    ]
    t3 = Table(disc_data, colWidths=[2*inch, 1.2*inch, 1.2*inch, 1*inch])
    t3.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
    ]))
    story.append(t3)
    story.append(Spacer(1, 16))

    # Student voices
    story.append(Paragraph("6. Student Voice (Selected Quotes)", h2))
    quotes = [
        '"I like it here because the teachers actually care about us. They know our names and ask how we\'re doing." — 5th grader',
        '"Some kids act up in the hallways but it\'s way better than last year. Mr. Chen is really good at talking to people." — 7th grader',
        '"I feel safe at school. If someone is being mean, I know I can talk to an adult and they\'ll actually do something." — 3rd grader',
        '"I wish we had more choices in what we learn about. My old school had genius hour and I miss that." — 6th grader',
        '"The library is my favorite place. Ms. Johnson always helps me find good books." — 4th grader',
    ]
    for q in quotes:
        story.append(Paragraph(f"<i>{q}</i>", body))
        story.append(Spacer(1, 8))

    # Overall assessment
    story.append(Spacer(1, 12))
    story.append(Paragraph("7. Overall Assessment", h2))
    story.append(Paragraph(
        "Lincoln Innovation Academy demonstrates a school culture that is meaningfully improving under "
        "current leadership. The investment in restorative practices, positive behavior systems, and "
        "relationship-building is yielding measurable results in discipline data and observable "
        "improvements in school climate. Key strengths include strong adult-student relationships, "
        "improving attendance, and a visible commitment to the school's mission. Growth areas include "
        "expanding student voice and agency opportunities (especially in middle school), ensuring "
        "consistent implementation of restorative practices across all classrooms, and increasing "
        "the display and celebration of student work in grades 6-8.", body))

    doc.build(story)
    print("  ✓ student_culture_walkthrough.pdf")


# ─── Main ────────────────────────────────────────────────────────────

def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Generating sample data in {OUT_DIR}/\n")

    generate_teacher_retention()
    generate_student_achievement()
    generate_budget()
    generate_classroom_observation()
    generate_family_survey()
    generate_culture_walkthrough()
    generate_pd_calendar()
    generate_leadership_notes()

    print(f"\nDone! {len(list(OUT_DIR.iterdir()))} files generated.")


if __name__ == "__main__":
    main()
